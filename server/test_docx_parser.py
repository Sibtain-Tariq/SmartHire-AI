import io
import docx
from app.services.resume_parser.docx_parser import DocxParser, DocxParserError

def create_dummy_docx(has_table: bool = False, empty: bool = False) -> bytes:
    doc = docx.Document()
    
    if not empty:
        doc.add_paragraph("John Doe - Resume")
        doc.add_paragraph("Software Engineer with 5 years of experience.")
        doc.add_paragraph("• Python\n• React\n• Node.js", style='List Bullet')
        
        if has_table:
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Degree"
            table.cell(0, 1).text = "University"
            table.cell(1, 0).text = "BS Computer Science"
            table.cell(1, 1).text = "MIT"
            
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def test_docx_parser():
    parser = DocxParser()
    
    print("--- Test 1: Normal DOCX ---")
    try:
        docx_bytes = create_dummy_docx()
        text = parser.extract_text(docx_bytes)
        print(f"Extracted:\n{text}")
        assert "John Doe - Resume" in text
        assert "Python" in text
        print("PASS")
    except Exception as e:
        print(f"FAIL: {e}")

    print("\n--- Test 2: DOCX with Tables ---")
    try:
        docx_bytes = create_dummy_docx(has_table=True)
        text = parser.extract_text(docx_bytes)
        print(f"Extracted:\n{text}")
        assert "BS Computer Science | MIT" in text
        print("PASS")
    except Exception as e:
        print(f"FAIL: {e}")

    print("\n--- Test 3: Invalid DOCX (Corrupted bytes) ---")
    try:
        invalid_bytes = b"This is not a word document"
        parser.extract_text(invalid_bytes)
        print("FAIL: Should have thrown an error")
    except DocxParserError as e:
        print(f"PASS: Caught expected error -> {e}")

    print("\n--- Test 4: Empty DOCX ---")
    try:
        empty_bytes = create_dummy_docx(empty=True)
        parser.extract_text(empty_bytes)
        print("FAIL: Should have thrown an error")
    except DocxParserError as e:
        print(f"PASS: Caught expected error -> {e}")

if __name__ == "__main__":
    test_docx_parser()
